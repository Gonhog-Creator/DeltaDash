"""
Pliego Técnico Matcher: Extracts requirements from uploaded bid/contract documents
and matches them against certified vest models in the database.

Uses Google Gemini 1.5 Flash (free tier) for structured requirement extraction,
then scores vests against the extracted requirements — only considering vests
that have official certifications (TestSession.is_official=True).
"""
import os
import json
import re
import logging
from typing import Dict, List, Optional, Any

from sqlalchemy.orm import Session

from app.core.config import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_file(file_path: str) -> str:
    """Extract text from a PDF or DOCX file.

    Returns the concatenated text content of the document.
    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT")


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ValueError(
            "PyMuPDF is not installed. Run: pip install PyMuPDF"
        )

    text_parts = []
    doc = fitz.open(file_path)
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts).strip()


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise ValueError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = docx.Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts).strip()


# ---------------------------------------------------------------------------
# Gemini AI Requirement Extraction
# ---------------------------------------------------------------------------

GEMINI_PROMPT = """\
You are an expert analyst for ballistic vest procurement contracts (pliegos técnicos).
Analyze the following document text and extract all ballistic vest requirements as structured JSON.

Return ONLY a JSON object with these fields (omit fields that are not mentioned in the document):

{
  "threat_level": "NIJ IIIA" | "NIJ III" | "NIJ II" | "Level IIA" | "RB1" | "RB2" | "RB3" | "RA1" | "RA2" | "RA3" | null,
  "protection_class": string | null,
  "vest_type": "soft" | "hard" | "tactical" | null,
  "required_sizes": ["S", "M", "L", "XL", "XXL"] | [],
  "max_weight_g": number | null,
  "trauma_homologation": {
    "backface_max_mm": number | null,
    "ammunition": string | null
  } | null,
  "flexibility_required": boolean | null,
  "panel_sewn_required": boolean | null,
  "ammunition_calibers": [".44 MAG", "9mm FMJ", "357 SIG", ...] | [],
  "is_female_required": boolean | null,
  "min_total_layers": number | null,
  "max_total_layers": number | null,
  "max_thickness_mm": number | null,
  "stitch_pattern": string | null,
  "additional_notes": string | null,
  "raw_summary": "A 2-3 sentence summary of the key requirements in the document"
}

Rules:
- If a value is not explicitly stated, omit the field or set it to null.
- Normalize threat levels to standard notation (NIJ IIIA, NIJ III, etc.).
- Extract all mentioned ammunition calibers.
- For weight, convert to grams if another unit is used.
- For backface deformation, convert to millimeters.
- Be conservative: only extract what is clearly stated in the document.
- ALWAYS output all text fields in English, regardless of the source document language.
- Translate any non-English terms (e.g., Spanish "nivel de amenaza" → "threat level") to English.
- The "raw_summary" field must always be written in English.

Document text:
"""


def extract_requirements_with_gemini(text: str) -> Dict[str, Any]:
    """Send document text to Gemini 1.5 Flash and get structured requirements back.

    Returns a dict of extracted requirements.
    Raises ValueError if the API key is not set or the API call fails.
    """
    api_key = settings.GEMINI_API_KEY
    if not api_key:
        raise ValueError(
            "GEMINI_API_KEY is not set. Add it to your .env file. "
            "Get a free key at https://aistudio.google.com/app/apikey"
        )

    try:
        import google.generativeai as genai
    except ImportError:
        raise ValueError(
            "google-generativeai is not installed. Run: pip install google-generativeai"
        )

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-3.6-flash")

    # Truncate text to avoid token limits (Gemini 1.5 Flash: ~1M tokens, but be safe)
    max_chars = 500_000
    truncated_text = text[:max_chars]
    if len(text) > max_chars:
        truncated_text += "\n\n[... document truncated ...]"

    full_prompt = GEMINI_PROMPT + truncated_text

    try:
        response = model.generate_content(
            full_prompt,
            generation_config={
                "temperature": 0.1,
                "top_p": 0.8,
                "max_output_tokens": 8192,
            },
        )

        raw_response = response.text.strip()

        # Strip markdown code fences if present
        if raw_response.startswith("```"):
            raw_response = re.sub(r"^```(?:json)?\s*", "", raw_response)
            raw_response = re.sub(r"\s*```$", "", raw_response)

        # Try to extract JSON even if there's extra text around it
        json_start = raw_response.find("{")
        json_end = raw_response.rfind("}")
        if json_start != -1 and json_end != -1 and json_end > json_start:
            json_str = raw_response[json_start:json_end + 1]
        else:
            json_str = raw_response

        requirements = json.loads(json_str)
        logger.info(f"Gemini extracted requirements: {requirements}")
        return requirements

    except json.JSONDecodeError as e:
        logger.error(f"Gemini returned invalid JSON: {e}")
        raise ValueError(f"AI returned invalid JSON. Raw response: {raw_response[:500]}")
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        raise ValueError(f"AI extraction failed: {str(e)}")


# ---------------------------------------------------------------------------
# Vest Matching
# ---------------------------------------------------------------------------

def match_vests(requirements: Dict[str, Any], db: Session) -> Dict[str, Any]:
    """Match extracted requirements against certified vests in the database.

    Only considers vests that have at least one TestSession with is_official=True.
    Returns a dict with:
        - recommendations: sorted list of vest matches with scores
        - summary: stats about the matching process
        - gaps: list of requirements that no vest satisfies
    """
    from app.db.models import Vest, VestLayer, Material
    from app.db.models.test_session import TestSession

    # Find vests with official certifications
    official_vest_ids = (
        db.query(TestSession.vest_id)
        .filter(TestSession.is_official.is_(True))
        .filter(TestSession.vest_id.isnot(None))
        .distinct()
        .all()
    )
    official_vest_id_set = {str(row[0]) for row in official_vest_ids}

    if not official_vest_id_set:
        return {
            "recommendations": [],
            "summary": {
                "total_certified_vests": 0,
                "total_matched": 0,
            },
            "gaps": ["No vests with official certifications found in the database."],
        }

    # Fetch those vests, optionally filtering by catalog models
    query = db.query(Vest).filter(Vest.id.in_(list(official_vest_id_set)))

    # Apply filters from requirements
    req_threat = requirements.get("threat_level")
    if req_threat:
        query = query.filter(Vest.threat_level.ilike(f"%{req_threat}%"))

    req_vest_type = requirements.get("vest_type")
    if req_vest_type:
        query = query.filter(Vest.vest_type.ilike(f"%{req_vest_type}%"))

    vests = query.all()

    scored = []
    for vest in vests:
        score, details, gaps = _score_vest(vest, requirements, db)
        if score > 0:
            # Fetch certification info
            cert_sessions = (
                db.query(TestSession)
                .filter(TestSession.vest_id == vest.id, TestSession.is_official.is_(True))
                .all()
            )
            certifications = [
                {
                    "name": s.name,
                    "lab_name": s.lab_name,
                    "protocol": s.protocol,
                    "certification_number": s.certification_number,
                    "test_date": s.test_date.isoformat() if s.test_date else None,
                }
                for s in cert_sessions
            ]

            scored.append({
                "vest_id": str(vest.id),
                "vest_code": vest.vest_code,
                "vest_type": vest.vest_type,
                "threat_level": vest.threat_level,
                "protection_class": vest.protection_class,
                "total_layers": vest.total_layers,
                "total_thickness_mm": float(vest.total_thickness_mm) if vest.total_thickness_mm else None,
                "weight_g": float(vest.weight_g) if vest.weight_g else None,
                "sizes": vest.sizes,
                "composition": vest.composition,
                "flexibility_rating": vest.flexibility_rating,
                "is_panel_sewn": vest.is_panel_sewn,
                "is_catalog_model": vest.is_catalog_model,
                "is_female": vest.is_female,
                "certifications": certifications,
                "match_score": round(score, 1),
                "match_details": details,
                "match_gaps": gaps,
            })

    # Sort by match score descending
    scored.sort(key=lambda x: x["match_score"], reverse=True)

    # Compute global gaps
    global_gaps = _compute_global_gaps(requirements, scored)

    return {
        "recommendations": scored,
        "summary": {
            "total_certified_vests": len(official_vest_id_set),
            "total_matched": len(scored),
            "top_score": scored[0]["match_score"] if scored else 0,
        },
        "gaps": global_gaps,
    }


def _score_vest(
    vest, requirements: Dict[str, Any], db: Session
) -> tuple:
    """Score a single vest against the extracted requirements.

    Returns (score, details_dict, gaps_list).
    Score is 0-100.
    """
    score = 0.0
    max_score = 100.0
    details = {}
    gaps = []

    # 1. Threat level match (30 points)
    req_threat = requirements.get("threat_level")
    if req_threat:
        vest_threat = (vest.threat_level or "").upper()
        req_threat_upper = req_threat.upper()
        # Normalize for comparison
        if req_threat_upper in vest_threat or vest_threat in req_threat_upper:
            score += 30
            details["threat_level"] = "match"
        else:
            gaps.append(f"Threat level: requires {req_threat}, vest has {vest.threat_level}")
            details["threat_level"] = "mismatch"
    else:
        score += 15  # Partial credit if no requirement specified
        details["threat_level"] = "not_specified"

    # 2. Weight compliance (20 points)
    max_weight = requirements.get("max_weight_g")
    if max_weight and vest.weight_g:
        vest_weight = float(vest.weight_g)
        if vest_weight <= max_weight:
            score += 20
            details["weight"] = "compliant"
        else:
            excess = ((vest_weight - max_weight) / max_weight) * 100
            if excess < 10:
                score += 10
                details["weight"] = "marginal"
            else:
                score += 0
                details["weight"] = "exceeds"
            gaps.append(f"Weight: {vest_weight}g exceeds max {max_weight}g by {excess:.1f}%")
    elif max_weight and not vest.weight_g:
        score += 10  # Unknown weight, partial credit
        details["weight"] = "unknown"
    else:
        score += 10
        details["weight"] = "not_specified"

    # 3. Trauma homologation (20 points)
    req_trauma = requirements.get("trauma_homologation")
    if req_trauma:
        req_backface = req_trauma.get("backface_max_mm")
        vest_trauma = vest.trauma_homologation
        if req_backface and vest_trauma:
            vest_backface = vest_trauma.get("backface_mm")
            if vest_backface and float(vest_backface) <= float(req_backface):
                score += 20
                details["trauma"] = "compliant"
            else:
                score += 5
                details["trauma"] = "marginal"
                gaps.append(f"Backface deformation: {vest_backface}mm vs max {req_backface}mm")
        else:
            score += 8
            details["trauma"] = "unknown"
    else:
        score += 10
        details["trauma"] = "not_specified"

    # 4. Size availability (15 points)
    req_sizes = requirements.get("required_sizes", [])
    if req_sizes:
        vest_sizes = vest.sizes or {}
        vest_size_keys = set((vest_sizes.keys() if isinstance(vest_sizes, dict) else []))
        req_size_set = set(req_sizes)
        matching = req_size_set & vest_size_keys
        if matching:
            score += 15 * (len(matching) / len(req_size_set))
            details["sizes"] = f"{len(matching)}/{len(req_size_set)} matched"
            if len(matching) < len(req_size_set):
                missing = req_size_set - vest_size_keys
                gaps.append(f"Missing sizes: {', '.join(sorted(missing))}")
        else:
            score += 0
            details["sizes"] = "no match"
            gaps.append(f"No required sizes available (needed: {', '.join(sorted(req_size_set))})")
    else:
        score += 7
        details["sizes"] = "not_specified"

    # 5. Flexibility (5 points)
    req_flex = requirements.get("flexibility_required")
    if req_flex is not None:
        if vest.flexibility_rating == req_flex:
            score += 5
            details["flexibility"] = "match"
        else:
            gaps.append(f"Flexibility: requires {req_flex}, vest has {vest.flexibility_rating}")
            details["flexibility"] = "mismatch"
    else:
        score += 2
        details["flexibility"] = "not_specified"

    # 6. Panel sewn (5 points)
    req_panel = requirements.get("panel_sewn_required")
    if req_panel is not None:
        if vest.is_panel_sewn == req_panel:
            score += 5
            details["panel_sewn"] = "match"
        else:
            gaps.append(f"Panel sewn: requires {req_panel}, vest has {vest.is_panel_sewn}")
            details["panel_sewn"] = "mismatch"
    else:
        score += 2
        details["panel_sewn"] = "not_specified"

    # 7. Catalog model bonus (5 points)
    if vest.is_catalog_model:
        score += 5
        details["catalog_model"] = True
    else:
        details["catalog_model"] = False

    # Clamp score
    score = min(score, max_score)

    return score, details, gaps


def _compute_global_gaps(requirements: Dict[str, Any], scored: List[Dict]) -> List[str]:
    """Identify requirements that no vest satisfies."""
    gaps = []

    if not scored:
        return ["No vests matched the extracted requirements."]

    # Check if any vest has a perfect threat level match
    req_threat = requirements.get("threat_level")
    if req_threat:
        any_match = any(d["match_details"].get("threat_level") == "match" for d in scored)
        if not any_match:
            gaps.append(f"No certified vest matches threat level '{req_threat}'.")

    # Check weight
    max_weight = requirements.get("max_weight_g")
    if max_weight:
        any_compliant = any(
            d["match_details"].get("weight") == "compliant" for d in scored
        )
        if not any_compliant:
            gaps.append(f"No vest meets the maximum weight requirement of {max_weight}g.")

    # Check sizes
    req_sizes = requirements.get("required_sizes", [])
    if req_sizes:
        any_full_size_match = any(
            "matched" in d["match_details"].get("sizes", "")
            and d["match_details"]["sizes"].startswith(f"{len(req_sizes)}/{len(req_sizes)}")
            for d in scored
        )
        if not any_full_size_match:
            gaps.append(f"No vest covers all required sizes: {', '.join(req_sizes)}.")

    # Check trauma
    req_trauma = requirements.get("trauma_homologation")
    if req_trauma and req_trauma.get("backface_max_mm"):
        any_compliant = any(
            d["match_details"].get("trauma") == "compliant" for d in scored
        )
        if not any_compliant:
            gaps.append(
                f"No vest meets the backface deformation limit of {req_trauma['backface_max_mm']}mm."
            )

    return gaps


# ---------------------------------------------------------------------------
# Full Pipeline
# ---------------------------------------------------------------------------

def analyze_pliego(file_path: str, db: Session) -> Dict[str, Any]:
    """Full pipeline: extract text → extract requirements via Gemini → match vests.

    Returns a dict with:
        - requirements: extracted requirements dict
        - match_results: vest matching results
    """
    # Step 1: Extract text
    text = extract_text_from_file(file_path)

    if not text or len(text.strip()) < 10:
        raise ValueError(
            "No readable text could be extracted from the document. "
            "If it's a scanned PDF, ensure it has a text layer."
        )

    # Step 2: Extract requirements via Gemini
    requirements = extract_requirements_with_gemini(text)

    # Step 3: Match against certified vests
    match_results = match_vests(requirements, db)

    return {
        "requirements": requirements,
        "match_results": match_results,
    }
